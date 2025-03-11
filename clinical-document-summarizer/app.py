import streamlit as st
import anthropic
import os
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from fpdf import FPDF
import json
import tempfile
from fpdf.enums import XPos, YPos
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Try to load from .env file for local development
load_dotenv()

def get_api_key():
    """Get API key from environment variables or Streamlit secrets"""
    # First try to get from Streamlit secrets (for cloud deployment)
    try:
        return st.secrets["ANTHROPIC_API_KEY"]
    except:
        # If not in secrets, try environment variable (for local development)
        return os.getenv('ANTHROPIC_API_KEY', '')

def initialize_client(api_key):
    """Initialize Anthropic client with the provided API key"""
    if api_key:
        return anthropic.Client(api_key=api_key)
    return None

# Define the Claude model to use
CLAUDE_MODEL = "claude-3-5-haiku-20241022"  # Update this when newer versions are available

# Initialize client variable
client = None

def extract_text_from_pdf(pdf_file):
    """Extract text from uploaded PDF file with better formatting"""
    pdf_reader = PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text

def parse_document_with_claude(text):
    """Use Claude to parse and structure the document with improved prompting"""
    prompt = f"""
    Analyze this clinical protocol document and extract key information in JSON format.
    Return ONLY a valid JSON object with the following structure, no other text:
    {{
        "study_title": "string",
        "clinical_phase": "string",
        "study_objectives": {{
            "primary": "string",
            "secondary": "string",
            "exploratory": "string"
        }},
        "study_rationale": "string",
        "study_population": "string",
        "inclusion_exclusion_criteria": "string",
        "primary_endpoints": "string",
        "secondary_exploratory_endpoints": "string",
        "study_design": "string",
        "subject_number": "string",
        "treatment_duration": "string",
        "duration_of_follow_up": "string",
        "dose_levels": "string",
        "route_of_delivery": "string",
        "data_safety_monitoring": "string",
        "stopping_rules": "string",
        "immune_monitoring": "string",
        "supporting_studies": "string",
        "assays_methodologies": "string",
        "statistical_analysis": "string",
        "outcome_criteria": "string",
        "risks": "string",
        "clinical_sites": "string",
        "clinical_operations": "string",
        "enrollment": "string",
        "long_term_follow_up": "string",
        "timeline": "string"
    }}

    Use "Not specified" if information is not found.
    Format lists as comma-separated strings.
    Avoid newlines in values.

    Document text:
    {text}
    """
    
    try:
        message = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=4000,
            temperature=0.0,
            system="You are a clinical protocol analyzer. Return only valid JSON without any markdown formatting or additional text.",
            messages=[{
                "role": "user",
                "content": prompt
            }]
        )
        
        # Get the response text
        if not message.content or not message.content[0].text:
            raise ValueError("Empty response from Claude")
            
        response_text = message.content[0].text.strip()
        
        # Remove any markdown formatting if present
        if response_text.startswith("```"):
            response_text = response_text.split("```")[1]
            if response_text.startswith("json"):
                response_text = response_text[4:]
        response_text = response_text.strip()
        
        # Parse the JSON response
        try:
            parsed_data = json.loads(response_text)
        except json.JSONDecodeError:
            st.error("Failed to parse JSON response. Raw response:")
            st.code(response_text)
            raise
        
        # Validate required keys
        required_keys = [
            "study_title", "clinical_phase", "study_objectives", "study_rationale",
            "study_population", "inclusion_exclusion_criteria", "primary_endpoints",
            "secondary_exploratory_endpoints", "study_design", "subject_number",
            "treatment_duration", "duration_of_follow_up", "dose_levels",
            "route_of_delivery", "data_safety_monitoring", "stopping_rules",
            "immune_monitoring", "supporting_studies", "assays_methodologies",
            "statistical_analysis", "outcome_criteria", "risks", "clinical_sites",
            "clinical_operations", "enrollment", "long_term_follow_up", "timeline"
        ]
        
        # Ensure all required keys exist
        for key in required_keys:
            if key not in parsed_data:
                parsed_data[key] = "Not specified"
            elif not parsed_data[key]:
                parsed_data[key] = "Not specified"
        
        return parsed_data
        
    except json.JSONDecodeError as e:
        st.error(f"Error parsing JSON response: {str(e)}")
        return None
    except Exception as e:
        st.error(f"An error occurred while parsing the document: {str(e)}")
        return None

def summarize_with_claude(section_content, section_name):
    """Use Claude to generate concise summaries with context-aware prompting"""
    if not section_content or section_content == "Not specified":
        return "Not specified"
        
    prompt = f"""
    Create a clear and concise summary of this {section_name} section from a clinical protocol.
    
    Guidelines:
    - Keep critical information
    - Use professional medical language
    - Format in clear paragraphs
    - Maintain list format for criteria
    - Be concise but complete
    
    Content to summarize:
    {section_content}
    """
    
    try:
        message = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=2000,
            temperature=0.3,
            system="You are a clinical protocol summarizer. Create clear, professional summaries.",
            messages=[{
                "role": "user",
                "content": prompt
            }]
        )
        
        if not message.content or not message.content[0].text:
            return "Error: No summary generated"
            
        summary = message.content[0].text.strip()
        return summary if summary else "Not specified"
        
    except Exception as e:
        st.error(f"Error generating summary for {section_name}: {str(e)}")
        return "Error generating summary"

def clean_text_for_pdf(text):
    """Clean text to handle special characters"""
    if not text:
        return "Not specified"
    # Replace problematic characters with their closest ASCII equivalents
    replacements = {
        '≥': '>=',
        '≤': '<=',
        '±': '+/-',
        '×': 'x',
        '°': 'deg',
        '–': '-',
        '—': '-',
        '"': '"',
        '"': '"',
        ''': "'",
        ''': "'",
        '…': '...',
        '•': '*'
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text

def generate_pdf_summary(data):
    """Generate PDF summary with better space handling"""
    class PDF(FPDF):
        def header(self):
            self.set_font('Helvetica', 'B', 16)
            self.cell(0, 10, "CLINICAL PROTOCOL SYNOPSIS", new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')
            self.ln(10)

    # Create PDF with A4 format
    pdf = PDF(format='A4')
    
    # Set margins (left, top, right)
    pdf.set_margins(25, 25, 25)
    
    # Set font
    pdf.set_font('Helvetica', size=12)
    
    # Add first page
    pdf.add_page()
    
    # Enable auto page break
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # Protocol Information Section
    sections = [
        ("STUDY TITLE", data["study_title"]),
        ("CLINICAL PHASE", data["clinical_phase"]),
        ("STUDY OBJECTIVES", f"Primary Objectives:\n{data['study_objectives']['primary']}\n\nSecondary Objectives:\n{data['study_objectives']['secondary']}\n\nExploratory Objectives:\n{data['study_objectives']['exploratory']}"),
        ("STUDY RATIONALE", data["study_rationale"]),
        ("STUDY POPULATION", data["study_population"]),
        ("MAIN INCLUSION/EXCLUSION CRITERIA", data["inclusion_exclusion_criteria"]),
        ("PRIMARY ENDPOINT(S)", data["primary_endpoints"]),
        ("SECONDARY & EXPLORATORY ENDPOINTS", data["secondary_exploratory_endpoints"]),
        ("STUDY DESIGN", data["study_design"]),
        ("SUBJECT NUMBER", data["subject_number"]),
        ("TREATMENT DURATION", data["treatment_duration"]),
        ("DURATION OF FOLLOW UP", data["duration_of_follow_up"]),
        ("DOSE LEVEL(S) AND DOSE JUSTIFICATION", data["dose_levels"]),
        ("ROUTE OF DELIVERY", data["route_of_delivery"]),
        ("DATA AND SAFETY MONITORING PLAN (DSMP)", data["data_safety_monitoring"]),
        ("STOPPING RULES", data["stopping_rules"]),
        ("IMMUNE MONITORING & IMMUNOSUPPRESSION", data["immune_monitoring"]),
        ("SUPPORTING STUDIES", data["supporting_studies"]),
        ("ASSAYS/METHODOLOGIES", data["assays_methodologies"]),
        ("STATISTICAL ANALYSIS PLAN", data["statistical_analysis"]),
        ("OUTCOME CRITERIA", data["outcome_criteria"]),
        ("RISKS", data["risks"]),
        ("CLINICAL SITES", data["clinical_sites"]),
        ("CLINICAL OPERATIONS", data["clinical_operations"]),
        ("ENROLLMENT", data["enrollment"]),
        ("LONG TERM FOLLOW UP", data["long_term_follow_up"]),
        ("TIMELINE", data["timeline"])
    ]
    
    for section_title, content in sections:
        # Section Title
        pdf.set_font('Helvetica', 'B', 12)
        pdf.cell(0, 10, section_title, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        
        # Section Content - clean the text before adding to PDF
        pdf.set_font('Helvetica', '', 11)
        cleaned_content = clean_text_for_pdf(content)
        
        # Split content into paragraphs
        paragraphs = cleaned_content.split('\n')
        
        for paragraph in paragraphs:
            if paragraph.strip():
                # Handle long paragraphs by breaking them into smaller chunks
                words = paragraph.split()
                current_line = []
                for word in words:
                    current_line.append(word)
                    if len(' '.join(current_line)) > 60:  # Further reduced line length
                        pdf.multi_cell(0, 8, ' '.join(current_line[:-1]))
                        current_line = [word]
                if current_line:
                    pdf.multi_cell(0, 8, ' '.join(current_line))
                pdf.ln(5)
        
        pdf.ln(5)
    
    return pdf

def generate_docx_summary(data):
    """Generate DOCX summary"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('CLINICAL PROTOCOL SYNOPSIS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sections = [
        ("STUDY TITLE", data["study_title"]),
        ("CLINICAL PHASE", data["clinical_phase"]),
        ("STUDY OBJECTIVES", data["study_objectives"]),
        ("STUDY RATIONALE", data["study_rationale"]),
        ("STUDY POPULATION", data["study_population"]),
        ("MAIN INCLUSION/EXCLUSION CRITERIA", data["inclusion_exclusion_criteria"]),
        ("PRIMARY ENDPOINT(S)", data["primary_endpoints"]),
        ("SECONDARY & EXPLORATORY ENDPOINTS", data["secondary_exploratory_endpoints"]),
        ("STUDY DESIGN", data["study_design"]),
        ("SUBJECT NUMBER", data["subject_number"]),
        ("TREATMENT DURATION", data["treatment_duration"]),
        ("DURATION OF FOLLOW UP", data["duration_of_follow_up"]),
        ("DOSE LEVEL(S) AND DOSE JUSTIFICATION", data["dose_levels"]),
        ("ROUTE OF DELIVERY", data["route_of_delivery"]),
        ("DATA AND SAFETY MONITORING PLAN (DSMP)", data["data_safety_monitoring"]),
        ("STOPPING RULES", data["stopping_rules"]),
        ("IMMUNE MONITORING & IMMUNOSUPPRESSION", data["immune_monitoring"]),
        ("SUPPORTING STUDIES", data["supporting_studies"]),
        ("ASSAYS/METHODOLOGIES", data["assays_methodologies"]),
        ("STATISTICAL ANALYSIS PLAN", data["statistical_analysis"]),
        ("OUTCOME CRITERIA", data["outcome_criteria"]),
        ("RISKS", data["risks"]),
        ("CLINICAL SITES", data["clinical_sites"]),
        ("CLINICAL OPERATIONS", data["clinical_operations"]),
        ("ENROLLMENT", data["enrollment"]),
        ("LONG TERM FOLLOW UP", data["long_term_follow_up"]),
        ("TIMELINE", data["timeline"])
    ]
    
    for section_title, content in sections:
        # Add section heading
        heading = doc.add_heading(section_title, level=1)
        heading.style.font.size = Pt(12)
        heading.style.font.bold = True
        
        # Add content
        paragraph = doc.add_paragraph()
        paragraph.add_run(content if content else "Not specified")
        
        # Add spacing after section
        doc.add_paragraph()
    
    return doc

def main():
    st.set_page_config(page_title="Clinical Protocol Summary Generator", layout="wide")
    
    st.title("Clinical Protocol Summary Generator")
    st.write("Upload a clinical study protocol document to generate a summarized version")

    # Initialize data variable
    data = None
    
    # Move file upload to sidebar
    with st.sidebar:
        st.header("Configuration")
        
        # Get initial API key
        initial_api_key = get_api_key()
        
        # API Key input field
        api_key = st.text_input(
            "Anthropic API Key", 
            type="password",
            value=initial_api_key,
            help="Enter your Anthropic API key. For Streamlit Cloud deployment, set this in your app's secrets."
        )

        # Initialize or update client
        if api_key:
            try:
                global client
                client = initialize_client(api_key)
                if client:
                    st.success("✅ API Key configured successfully!")
                else:
                    st.error("❌ Invalid API Key format")
                    st.stop()
            except Exception as e:
                st.error(f"❌ Error initializing client: {str(e)}")
                st.stop()
        else:
            st.error("❌ Please provide your Anthropic API Key")
            st.info("💡 You can get an API key from [Anthropic Console](https://console.anthropic.com)")
            st.info("📝 For Streamlit Cloud: Add your API key in Settings → Secrets")
            st.stop()

        st.header("Upload Protocol")
        uploaded_file = st.file_uploader("Choose a PDF file", type=['pdf'])
        
        if uploaded_file is not None:
            if st.button("Process Document"):
                with st.spinner("Processing document..."):
                    # Extract text from PDF
                    text = extract_text_from_pdf(uploaded_file)
                    
                    # Parse document structure
                    with st.spinner("Analyzing document structure..."):
                        data = parse_document_with_claude(text)
                    
                    if data:
                        st.success("Document processed successfully!")
                        
                        # Generate DOCX summary
                        doc = generate_docx_summary(data)
                        
                        # Save the document
                        doc_path = "protocol_summary.docx"
                        doc.save(doc_path)
                        
                        # Provide download button
                        with open(doc_path, "rb") as file:
                            st.download_button(
                                label="Download Summary",
                                data=file,
                                file_name="protocol_summary.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        
                        # Store data in session state
                        st.session_state['protocol_data'] = data
                    else:
                        st.error("Failed to parse document structure. Please try again.")

    # Main content area - only show if data is available
    if 'protocol_data' in st.session_state:
        data = st.session_state['protocol_data']
        st.header("Document Analysis")
        
        # Display sections in a grid layout
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Study Information")
            st.write("**Study Title:**", data.get("study_title", "Not specified"))
            st.write("**Clinical Phase:**", data.get("clinical_phase", "Not specified"))
            st.write("**Study Population:**", data.get("study_population", "Not specified"))
            st.write("**Subject Number:**", data.get("subject_number", "Not specified"))
            
            st.subheader("Study Design")
            st.write("**Study Design:**", data.get("study_design", "Not specified"))
            st.write("**Treatment Duration:**", data.get("treatment_duration", "Not specified"))
            st.write("**Follow-up Duration:**", data.get("duration_of_follow_up", "Not specified"))
            st.write("**Route of Delivery:**", data.get("route_of_delivery", "Not specified"))
            
            st.subheader("Endpoints")
            st.write("**Primary Endpoints:**", data.get("primary_endpoints", "Not specified"))
            st.write("**Secondary and Exploratory Endpoints:**", data.get("secondary_exploratory_endpoints", "Not specified"))
            
            st.subheader("Safety and Monitoring")
            st.write("**Data and Safety Monitoring Plan:**", data.get("data_safety_monitoring", "Not specified"))
            st.write("**Stopping Rules:**", data.get("stopping_rules", "Not specified"))
            st.write("**Immune Monitoring and Immunosuppression:**", data.get("immune_monitoring", "Not specified"))
            st.write("**Risks:**", data.get("risks", "Not specified"))
        
        with col2:
            st.subheader("Objectives and Rationale")
            st.write("**Study Objectives:**", data.get("study_objectives", "Not specified"))
            st.write("**Study Rationale:**", data.get("study_rationale", "Not specified"))
            
            st.subheader("Criteria")
            st.write("**Main Inclusion/Exclusion Criteria:**", data.get("inclusion_exclusion_criteria", "Not specified"))
            st.write("**Outcome Criteria:**", data.get("outcome_criteria", "Not specified"))
            
            st.subheader("Dosing and Analysis")
            st.write("**Dose Levels and Dose Justification:**", data.get("dose_levels", "Not specified"))
            st.write("**Statistical Analysis Plan:**", data.get("statistical_analysis", "Not specified"))
            
            st.subheader("Operations")
            st.write("**Clinical Sites:**", data.get("clinical_sites", "Not specified"))
            st.write("**Clinical Operations Plan:**", data.get("clinical_operations", "Not specified"))
            st.write("**Enrollment:**", data.get("enrollment", "Not specified"))
            st.write("**Long Term Follow Up:**", data.get("long_term_follow_up", "Not specified"))
            st.write("**Timeline:**", data.get("timeline", "Not specified"))
            
            st.subheader("Supporting Information")
            st.write("**Supporting Studies:**", data.get("supporting_studies", "Not specified"))
            st.write("**Assays/Methodologies:**", data.get("assays_methodologies", "Not specified"))

if __name__ == "__main__":
    main() 